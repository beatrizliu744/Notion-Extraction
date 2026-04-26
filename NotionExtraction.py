# notion_export_slim.py
# -*- coding: utf-8 -*-
"""
Export three Notion databases into ONE compact DOCX for 2026-03-01 .. 2026-03-31
- Tasks:   table -> Task name | Due Date | Page Content  (Due Date in range & Status=Done)
- Mood:    per-page rich render (images supported)
- Learning:table -> Date | Category | Language | Content | Notes | Source | Tags | Related | Page Content
  (Page Content 为可全文检索的纯文本，包含内嵌表格文本/Toggle 摘要/图片说明等)

Deps: requests, python-docx  (auto安装)
"""

import os, sys, json, time, tempfile, subprocess, datetime as dt
from typing import Dict, Any, List, Optional, Tuple

# ========= Config =========
NOTION_TOKEN = ""  # 你的 Notion Integration Secret

YEAR = 2026
START = f"{YEAR}-03-01"
END   = f"{YEAR}-03-31"

DB_TASKS    = "24f5e9cc51e080b5b49ef7ab3232aa15"
DB_MOOD     = "24f5e9cc51e081c19306fd36e73e4f3e"
DB_LEARNING = "24f5e9cc51e080b8b573ed5ef3230328"

OUT_DOCX = os.path.join(os.path.expanduser("~"), "Desktop", "26_notion_0301-0331.docx")

# ========= Bootstrap deps =========
def _ensure_deps():
    try:
        import requests, docx  # noqa
    except Exception:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "requests", "python-docx"])
_ensure_deps()

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

NOTION_API = "https://api.notion.com/v1"
NOTION_VERSION = "2022-06-28"

# ========= Utils =========
def iso_date(s: str) -> dt.date:
    return dt.datetime.strptime(s, "%Y-%m-%d").date()

def overlap(a_start: Optional[dt.date], a_end: Optional[dt.date], b_start: dt.date, b_end: dt.date) -> bool:
    if a_start is None and a_end is None:
        return False
    if a_start is None:
        a_start = a_end
    if a_end is None:
        a_end = a_start
    return a_start <= b_end and a_end >= b_start

# ========= Notion Client =========
class NotionClient:
    def __init__(self) -> None:
        token = NOTION_TOKEN or os.getenv("NOTION_TOKEN")
        if not token:
            raise RuntimeError("No Notion token. Set NOTION_TOKEN in script or env.")
        self.sess = requests.Session()
        self.sess.headers.update({
            "Authorization": f"Bearer {token}",
            "Notion-Version": NOTION_VERSION,
            "Content-Type": "application/json; charset=utf-8",
        })

    def _req(self, method: str, path: str, **kw) -> Dict[str, Any]:
        url = f"{NOTION_API}{path}"
        for _ in range(5):
            r = self.sess.request(method, url, timeout=60, **kw)
            if r.status_code in (200, 201):
                return r.json()
            if r.status_code in (429, 500, 502, 503, 504):
                time.sleep(1.5)
                continue
            try:
                j = r.json()
            except Exception:
                j = {"error": r.text}
            raise RuntimeError(f"Notion API {r.status_code}: {json.dumps(j)}")
        raise RuntimeError("Notion API retries exceeded")

    def db(self, dbid: str) -> Dict[str, Any]:
        return self._req("GET", f"/databases/{dbid}")

    def page(self, pid: str) -> Dict[str, Any]:
        return self._req("GET", f"/pages/{pid}")

    def comments(self, block_id: str) -> List[Dict[str, Any]]:
        data = self._req("GET", "/comments", params={"block_id": block_id})
        return data.get("results", [])

    def date_prop(self, dbid: str, preferred: Optional[str]) -> Optional[str]:
        try:
            props = self.db(dbid).get("properties", {})
            if preferred in props and props[preferred].get("type") == "date":
                return preferred
            for name, meta in props.items():
                if meta.get("type") == "date" and name.lower() == "date":
                    return name
            for name, meta in props.items():
                if meta.get("type") == "date":
                    return name
        except Exception:
            pass
        return None

    def query(self, dbid: str, date_prop: Optional[str], end_date: dt.date) -> List[Dict[str, Any]]:
        payload: Dict[str, Any] = {"page_size": 100}
        if date_prop:
            payload["filter"] = {"property": date_prop, "date": {"on_or_before": end_date.isoformat()}}
            payload["sorts"] = [{"property": date_prop, "direction": "ascending"}]
        res, cursor = [], None
        while True:
            if cursor:
                payload["start_cursor"] = cursor
            data = self._req("POST", f"/databases/{dbid}/query", json=payload)
            res.extend(data.get("results", []))
            if not data.get("has_more"):
                break
            cursor = data.get("next_cursor")
        return res

    def children(self, block_id: str) -> List[Dict[str, Any]]:
        out, cursor = [], None
        while True:
            params = {"page_size": 100}
            if cursor:
                params["start_cursor"] = cursor
            data = self._req("GET", f"/blocks/{block_id}/children", params=params)
            out.extend(data.get("results", []))
            if not data.get("has_more"):
                break
            cursor = data.get("next_cursor")
        return out

# ========= Property Helpers =========
def render_rich_text(rt: List[Dict[str, Any]]) -> str:
    out: List[str] = []
    for x in rt:
        t = x.get("type")
        if t == "text":
            txt = x["text"].get("content", "")
            link = x["text"].get("link", {}).get("url") if x["text"].get("link") else None
            ann = x.get("annotations", {})
            if ann.get("code"):
                txt = f"`{txt}`"
            if ann.get("bold"):
                txt = f"**{txt}**"
            if ann.get("italic"):
                txt = f"*{txt}*"
            if ann.get("strikethrough"):
                txt = f"~~{txt}~~"
            if ann.get("underline"):
                txt = f"<u>{txt}</u>"
            if link:
                txt = f"{txt} ({link})"
            out.append(txt)
        elif t == "mention":
            m = x.get("mention", {})
            if "date" in m:
                out.append(m["date"].get("start", ""))
        elif t == "equation":
            out.append(f"$ {x['equation'].get('expression','')} $")
    return "".join(out)

def page_title(page: Dict[str, Any]) -> str:
    for _, prop in page.get("properties", {}).items():
        if prop.get("type") == "title":
            return render_rich_text(prop.get("title", [])) or "Untitled"
    return "Untitled"

def extract_date_range(page: Dict[str, Any], date_prop: str) -> Tuple[Optional[dt.date], Optional[dt.date]]:
    prop = page.get("properties", {}).get(date_prop, {})
    if prop.get("type") != "date":
        return (None, None)
    d = prop.get("date") or {}
    s, e = d.get("start"), d.get("end")

    def parse(v: Optional[str]) -> Optional[dt.date]:
        if not v:
            return None
        try:
            return dt.datetime.fromisoformat(v.replace("Z", "+00:00")).date()
        except Exception:
            return dt.datetime.strptime(v[:10], "%Y-%m-%d").date()

    ds = parse(s)
    de = parse(e) if e else ds
    return ds, de

ALIASES = {
    "title":   ["Title", "名称", "Name", "Task name", "Task Name", "任务", "任务名"],
    "date":    ["Date", "日期", "Due Date", "Due date", "Due"],
    "emotion": ["Emotion", "情绪", "Mood"],
    "main_task": ["Main Task", "主线任务", "Primary Task"],
    "body":    ["Body Sensations", "身体感受", "Somatic", "Body"],
    "thoughts": ["Thoughts", "想法", "Ideas"],
    "category": ["Category", "类别", "Type"],
    "language": ["Language", "语种", "Lang"],
    "content": ["Content", "内容", "正文"],
    "notes":   ["Notes", "备注", "注释"],
    "source":  ["Source", "Source Link", "来源", "来源链接", "URL", "Link"],
    "tags":    ["Tags", "标签", "主题标签"],
    "related": ["Related", "Relation", "关联页面", "关联"],
}

def get_prop(page: Dict[str, Any], aliases: List[str]) -> Optional[Dict[str, Any]]:
    props = page.get("properties", {})
    lower = {k.lower(): k for k in props}
    for a in aliases:
        k = lower.get(a.lower())
        if k:
            return props[k]
    for a in aliases:
        for name in props:
            if a.lower() == name.lower():
                return props[name]
    return None

def prop_text(prop: Optional[Dict[str, Any]], client: Optional[NotionClient] = None) -> str:
    if not prop:
        return ""
    t = prop.get("type")
    if t == "title":
        return render_rich_text(prop.get("title", []))
    if t == "rich_text":
        return render_rich_text(prop.get("rich_text", []))
    if t == "date":
        d = prop.get("date") or {}
        s = d.get("start") or ""
        e = d.get("end")
        return f"{s}–{e}" if e else s
    if t == "status":
        return (prop.get("status") or {}).get("name", "") or ""
    if t == "select":
        return (prop.get("select") or {}).get("name", "") or ""
    if t == "multi_select":
        return ", ".join([x.get("name", "") for x in (prop.get("multi_select") or []) if x.get("name")])
    if t == "url":
        return prop.get("url") or ""
    if t == "people":
        return ", ".join([p.get("name", "") or p.get("id", "") for p in (prop.get("people") or [])])
    if t == "relation":
        xs = prop.get("relation") or []
        names = []
        for r in xs:
            pid = r.get("id")
            if client:
                try:
                    names.append(page_title(client.page(pid)))
                except Exception:
                    names.append(pid or "")
            else:
                names.append(pid or "")
        return ", ".join(names)
    if t == "number":
        return str(prop.get("number") or "")
    if t == "checkbox":
        return "Yes" if prop.get("checkbox") else "No"
    return ""

# ========= Flatten for searchable text =========
def flatten_blocks_to_text(blocks: List[Dict[str, Any]], client: NotionClient) -> str:
    lines: List[str] = []

    def walk(blks: List[Dict[str, Any]], indent: int = 0):
        pad = "  " * indent
        for b in blks:
            t = b.get("type")
            obj = b.get(t, {})

            def rt() -> str:
                return render_rich_text(obj.get("rich_text", []))

            if t in (
                "paragraph",
                "heading_1",
                "heading_2",
                "heading_3",
                "quote",
                "callout",
                "bulleted_list_item",
                "numbered_list_item",
            ):
                txt = rt()
                if txt:
                    lines.append(f"{pad}{txt}")

            elif t == "to_do":
                txt = rt()
                chk = "☑" if obj.get("checked") else "☐"
                lines.append(f"{pad}{chk} {txt}")

            elif t == "code":
                code = "".join(
                    [x["text"]["content"] for x in obj.get("rich_text", []) if x.get("type") == "text"]
                )
                if code:
                    lines.append(pad + code)

            elif t == "bookmark":
                url = obj.get("url", "")
                if url:
                    lines.append(pad + f"[Link] {url}")
                else:
                    lines.append(pad + "[Link]")

            elif t == "image":
                cap = render_rich_text(obj.get("caption", []))
                if cap:
                    lines.append(pad + "[Image] " + cap)
                else:
                    lines.append(pad + "[Image]")

            elif t == "equation":
                lines.append(pad + f"$ {obj.get('expression', '')} $")

            elif t == "toggle":
                summary = rt()
                if summary:
                    lines.append(pad + summary)

            elif t == "table":
                rows = []
                for row in client.children(b["id"]):
                    if row.get("type") != "table_row":
                        continue
                    cells = row["table_row"]["cells"]
                    rows.append(" | ".join(render_rich_text(c) for c in cells))
                if rows:
                    lines.append(pad + "[Table]")
                    for r in rows:
                        lines.append(pad + "  " + r)

            # 递归子块（toggle 不增加缩进）
            if b.get("has_children"):
                walk(client.children(b["id"]), indent + (0 if t == "toggle" else 1))

    walk(blocks)
    return "\n".join([x for x in lines if x.strip()]).strip()

# ========= DOCX renderer =========
class DocxRenderer:
    def __init__(self, title: str) -> None:
        self.doc = Document()
        # 紧凑页边距+字号
        for s in self.doc.sections:
            s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.6)
        try:
            self.doc.styles["Normal"].font.size = Pt(10.5)
        except Exception:
            pass
        if title:
            h = self.doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 代码样式
        if "Code" not in self.doc.styles:
            st = self.doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            st.font.name = "Consolas"
            st.font.size = Pt(9)

    def add_h1(self, text: str):
        self.doc.add_heading(text, level=1)

    def add_h2(self, text: str):
        self.doc.add_heading(text, level=2)

    def p(self, text: str = ""):
        self.doc.add_paragraph(text)

    def bullet(self, text: str, level: int = 0):
        p = self.doc.add_paragraph(text, style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)

    def numbered(self, text: str, level: int = 0):
        p = self.doc.add_paragraph(text, style="List Number")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)

    def table(self, headers: List[str], rows: List[List[str]]):
        if not headers:
            return
        t = self.doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid"
        except Exception:
            pass
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for r in rows:
            cells = t.add_row().cells
            for i in range(len(headers)):
                cells[i].text = (r[i] if i < len(r) else "") or ""
        t.autofit = True

    def divider(self):
        self.p("—" * 20)

    def image_from_url(self, url: str, caption: str = ""):
        try:
            r = requests.get(url, timeout=30)
            r.raise_for_status()
            ctype = r.headers.get("Content-Type", "")
            ext = (
                ".jpg" if "jpeg" in ctype or "jpg" in ctype else
                ".gif" if "gif" in ctype else
                ".bmp" if "bmp" in ctype else
                ".tif" if "tiff" in ctype or "tif" in ctype else
                ".png"
            )
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(r.content)
                path = tmp.name
            self.doc.add_picture(path, width=Inches(5.0))
            if caption:
                p = self.doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                os.unlink(path)
            except Exception:
                pass
        except Exception:
            self.p(f"[Image failed: {url}]")

def blocks_to_docx(blocks: List[Dict[str, Any]], client: NotionClient, R: DocxRenderer):
    for b in blocks:
        t = b.get("type")
        obj = b.get(t, {})
        if t == "paragraph":
            R.p(render_rich_text(obj.get("rich_text", [])))
        elif t in ("heading_1", "heading_2", "heading_3"):
            txt = render_rich_text(obj.get("rich_text", []))
            if t == "heading_1":
                R.add_h2(txt)
            elif t == "heading_2":
                R.doc.add_heading(txt, level=3)
            else:
                R.doc.add_heading(txt, level=4)
        elif t == "quote":
            p = R.doc.add_paragraph(render_rich_text(obj.get("rich_text", [])))
            try:
                p.style = R.doc.styles["Intense Quote"]
            except Exception:
                pass
        elif t == "bulleted_list_item":
            R.bullet(render_rich_text(obj.get("rich_text", [])))
        elif t == "numbered_list_item":
            R.numbered(render_rich_text(obj.get("rich_text", [])))
        elif t == "to_do":
            R.p(("☑ " if obj.get("checked") else "☐ ") + render_rich_text(obj.get("rich_text", [])))
        elif t == "callout":
            emoji = (obj.get("icon") or {}).get("emoji", "💡")
            R.p(f"{emoji} {render_rich_text(obj.get('rich_text', []))}")
        elif t == "code":
            R.doc.add_paragraph(
                "".join([x["text"]["content"] for x in obj.get("rich_text", []) if x.get("type") == "text"]),
                style="Code",
            )
        elif t == "toggle":
            R.p("▼ " + render_rich_text(obj.get("rich_text", [])))
        elif t == "divider":
            R.divider()
        elif t == "bookmark":
            url = obj.get("url", "")
            cap = render_rich_text(obj.get("caption", []))
            R.p(f"{cap or 'Link'}: {url}")
        elif t == "image":
            img = obj.get("file") or obj.get("external") or {}
            url = img.get("url", "")
            cap = render_rich_text(obj.get("caption", []))
            if url:
                R.image_from_url(url, caption=cap)
        elif t == "equation":
            R.p(f"$$ {obj.get('expression', '')} $$")
        elif t == "table":
            rows: List[List[str]] = []
            for row in client.children(b["id"]):
                if row.get("type") != "table_row":
                    continue
                cells = row["table_row"]["cells"]
                rows.append([render_rich_text(c) for c in cells])
            if rows:
                R.table(rows[0], rows[1:])
        if b.get("has_children") and t != "table":
            blocks_to_docx(client.children(b["id"]), client, R)

# ========= Exporters =========
def matches_status(page: Dict[str, Any], want: Optional[str]) -> bool:
    if not want:
        return True
    want = want.casefold()
    for _, prop in page.get("properties", {}).items():
        tp = prop.get("type")
        if tp == "status":
            name = (prop.get("status") or {}).get("name", "")
            if name and name.casefold() == want:
                return True
        elif tp == "select":
            name = (prop.get("select") or {}).get("name", "")
            if name and name.casefold() == want:
                return True
    return False

def export_tasks(R: DocxRenderer, C: NotionClient, start: dt.date, end: dt.date) -> int:
    date_prop = C.date_prop(DB_TASKS, "Due Date") or "Due Date"
    pages = C.query(DB_TASKS, date_prop, end)
    rows, cnt = [], 0
    for p in pages:
        ds, de = extract_date_range(p, date_prop)
        if not overlap(ds, de, start, end):
            continue
        if not matches_status(p, "Done"):
            continue
        title = prop_text(get_prop(p, ALIASES["title"]))
        duedate = prop_text(get_prop(p, ["Due Date", "Due date", "Due", "Date"]))
        flat = flatten_blocks_to_text(C.children(p["id"]), C)
        rows.append([title, duedate, flat])
        cnt += 1
    R.add_h1("Tasks")
    R.table(["Task name", "Due Date", "Page Content"], rows)
    return cnt

def export_mood(R: DocxRenderer, C: NotionClient, start: dt.date, end: dt.date) -> int:
    date_prop = C.date_prop(DB_MOOD, "Date") or "Date"
    pages = C.query(DB_MOOD, date_prop, end)
    cnt = 0
    R.add_h1("Mood")
    for p in pages:
        ds, de = extract_date_range(p, date_prop)
        if not overlap(ds, de, start, end):
            continue
        title = prop_text(get_prop(p, ALIASES["title"])) or page_title(p)
        date_txt = prop_text(get_prop(p, ALIASES["date"]))
        emotion = prop_text(get_prop(p, ALIASES["emotion"]))
        mtask = prop_text(get_prop(p, ALIASES["main_task"]))
        body = prop_text(get_prop(p, ALIASES["body"]))
        thoughts = prop_text(get_prop(p, ALIASES["thoughts"]))
        R.add_h2(f"{title or 'Untitled'} — {date_txt}")
        if emotion:
            R.p(f"Emotion: {emotion}")
        if mtask:
            R.p(f"Main Task: {mtask}")
        if body:
            R.p(f"Body: {body}")
        if thoughts:
            R.p(f"Thoughts: {thoughts}")
        R.p("Content:")
        blocks_to_docx(C.children(p["id"]), C, R)
        cmts = []
        try:
            cmts = C.comments(p["id"])
        except Exception:
            cmts = []
        if cmts:
            R.p("Comments:")
            for c in cmts:
                R.p(f"- [{c.get('created_time', '')[:19]}] {render_rich_text(c.get('rich_text', []))}")
        R.divider()
        cnt += 1
    return cnt

def export_learning(R: DocxRenderer, C: NotionClient, start: dt.date, end: dt.date) -> int:
    date_prop = C.date_prop(DB_LEARNING, "Date") or "Date"
    pages = C.query(DB_LEARNING, date_prop, end)
    rows, cnt = [], 0
    for p in pages:
        ds, de = extract_date_range(p, date_prop)
        if not overlap(ds, de, start, end):
            continue
        date_tx = prop_text(get_prop(p, ALIASES["date"]))
        cate = prop_text(get_prop(p, ALIASES["category"]))
        lang = prop_text(get_prop(p, ALIASES["language"]))
        content = prop_text(get_prop(p, ALIASES["content"]))
        notes = prop_text(get_prop(p, ALIASES["notes"]))
        source = prop_text(get_prop(p, ALIASES["source"]))
        tags = prop_text(get_prop(p, ALIASES["tags"]))
        related = prop_text(get_prop(p, ALIASES["related"]), client=C)
        flat = flatten_blocks_to_text(C.children(p["id"]), C)
        rows.append([date_tx, cate, lang, content, notes, source, tags, related, flat])
        cnt += 1
    R.add_h1("Learning")
    R.table(["Date", "Category", "Language", "Content", "Notes", "Source", "Tags", "Related", "Page Content"], rows)
    return cnt

# ========= Main =========
def main():
    C = NotionClient()
    start, end = iso_date(START), iso_date(END)
    if end < start:
        raise SystemExit("END before START")
    R = DocxRenderer(title=f"Notion Export [{START} .. {END}]")

    total = 0
    total += export_tasks(R, C, start, end)
    total += export_mood(R, C, start, end)
    total += export_learning(R, C, start, end)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    R.doc.save(OUT_DOCX)
    print(f"Saved DOCX: {OUT_DOCX}")
    print(f"Total items: {total}")

if __name__ == "__main__":
    main()
